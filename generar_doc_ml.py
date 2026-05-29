"""
MILPÍN — Generador de Documento ML y Notebooks
Genera MILPIN_Documento_ML_Notebooks.docx con python-docx 1.2+
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Rutas ──────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
IMG_ML = os.path.join(BASE, "ML", "images")
IMG_EXP = os.path.join(BASE, "ML", "experiments", "ImagenesML")
OUT = os.path.join(BASE, "MILPIN_Documento_ML_Notebooks.docx")

# ── Paleta de colores ──────────────────────────────────────────────────────
GREEN_DARK  = RGBColor(0x1A, 0x47, 0x2A)
GREEN_MED   = RGBColor(0x2D, 0x6A, 0x4F)
GREEN_LIGHT = RGBColor(0x52, 0xB7, 0x88)
BLUE_ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GRAY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
RED_WARN    = RGBColor(0xC0, 0x00, 0x00)

# ── Helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Fondo de celda en tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "52B788")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def img(doc, folder, filename, caption_text=None, width_cm=14):
    fpath = os.path.join(folder, filename)
    if not os.path.exists(fpath):
        p = doc.add_paragraph(f"[Imagen no disponible: {filename}]")
        p.runs[0].italic = True
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return
    doc.add_picture(fpath, width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_text:
        cp = doc.add_paragraph(caption_text)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].italic = True
        cp.runs[0].font.color.rgb = GREEN_MED

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = GREEN_DARK
    p.runs[0].font.size = Pt(16)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = GREEN_MED
    p.runs[0].font.size = Pt(13)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = BLUE_ACCENT
    p.runs[0].font.size = Pt(11)
    return p

def para(doc, text, bold=False, italic=False, size=10.5, color=None, indent=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def result_box(doc, text):
    """Caja verde oscuro con texto blanco para resultados clave."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1A472A")
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def info_box(doc, text, color_hex="EBF5EB"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    doc.add_paragraph()

def add_table(doc, headers, rows_data, col_widths_cm=None):
    """Tabla con encabezado verde y filas alternadas."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezado
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, "2D6A4F")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Filas de datos
    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx + 1]
        bg = "F2F9F5" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)

    # Anchos de columnas
    if col_widths_cm:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])

    doc.add_paragraph()
    return tbl

# ── Documento ──────────────────────────────────────────────────────────────
doc = Document()

# Página A4 y márgenes 2cm
from docx.shared import Mm
section = doc.sections[0]
section.page_width  = Mm(210)
section.page_height = Mm(297)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# Fuente por defecto
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n")
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MILPÍN")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = GREEN_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Apoyo a Decisiones Agrícolas\nValle del Yaqui, Sonora — DR-041, Módulo 3")
r.font.size = Pt(14)
r.font.color.rgb = GREEN_MED

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─" * 60)
r.font.color.rgb = GREEN_LIGHT

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Módulo de Machine Learning y Análisis de Datos")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = BLUE_ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentación Técnica de Notebooks, Modelos y Resultados")
r.font.size = Pt(13)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Mayo 2026  ·  Versión 1.0")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 0 — INTRODUCCIÓN Y CONTEXTO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "0. Introducción y Contexto del Módulo ML")

para(doc, (
    "MILPÍN es un sistema de apoyo a decisiones (DSS) agrícola diseñado para optimizar el uso del agua de riego "
    "en el Distrito de Riego DR-041 (Valle del Yaqui, Sonora), con foco inicial en el Módulo 3. El objetivo central "
    "del sistema es reducir el consumo de agua de 8,000 m³/ha por ciclo agrícola a 6,000 m³/ha — un ahorro del 25% — "
    "sin comprometer el rendimiento de los cultivos."
))

para(doc, (
    "Este documento describe íntegramente el módulo de Machine Learning: los experimentos realizados, los modelos "
    "entrenados, los resultados obtenidos y la arquitectura de inferencia en producción. Cubre todos los notebooks "
    "del proyecto (archivos .ipynb), los scripts de entrenamiento e inferencia, el sistema de monitoreo de drift "
    "y el Feature Store."
))

h2(doc, "KPI Central y Justificación Económica")
para(doc, (
    "La reducción objetivo de 2,000 m³/ha/ciclo equivale a un ahorro de $3,360 MXN/ha/ciclo a la tarifa baseline "
    "de $1.68 MXN/m³ (CFE 9-CU, bombeo a 80 m de profundidad). Con parcelas promedio de 10-20 ha y múltiples ciclos "
    "anuales, el impacto económico es significativo para los agricultores del módulo."
))

result_box(doc, "KPI: 8,000 → 6,000 m³/ha/ciclo  |  Ahorro: $3,360 MXN/ha/ciclo  |  Tarifa: $1.68 MXN/m³")

h2(doc, "Estructura del Módulo ML")
para(doc, (
    "El módulo ML sigue una arquitectura MLOps explícita con separación de responsabilidades entre entrenamiento, "
    "inferencia y datos:"
))
bullet(doc, "ml/training/ — código de entrenamiento, nunca importa backend/")
bullet(doc, "ml/inference/ — wrappers de inferencia (singletons) consumidos por la API")
bullet(doc, "ml/monitoring/ — detección de drift (PSI, KS) y métricas compartidas")
bullet(doc, "ml/feature_store/ — definiciones YAML de features por dominio")
bullet(doc, "ml/experiments/ — notebooks exploratorios (EDA, experimentos)")
bullet(doc, "ml/configs/ — hiperparámetros declarativos en YAML")
bullet(doc, "ml/pipelines/ — Prefect flows para automatización (stubs, Fase C/D)")

para(doc, (
    "Los modelos de producción son: (1) XGBoost v4 con tres submodelos por cultivo para recomendación de riego, "
    "(2) Isolation Forest para detección de anomalías en patrones de riego, y (3) Ridge Regression para forecast "
    "de ETo a 7 días."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 1 — EDA
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "1. Análisis Exploratorio de Datos — eda_milpin.ipynb")

h2(doc, "1.1 Objetivo")
para(doc, (
    "El notebook de EDA establece el entendimiento del dominio antes de construir cualquier modelo. "
    "El objetivo es responder: ¿qué variables describen una parcela agrícola en el Valle del Yaqui?, "
    "¿cómo se distribuyen los volúmenes de riego históricamente?, y ¿cuál es la exposición económica "
    "al riego ineficiente?"
))
para(doc, (
    "El EDA no es un paso ceremonial — en MILPÍN define el Índice de Riesgo Hidráulico, que se convierte "
    "en feature de entrada para XGBoost y en métrica de priorización de parcelas."
))

h2(doc, "1.2 Qué se hizo")
para(doc, "El notebook carga 6 datasets del sistema y realiza:")
bullet(doc, "Análisis de distribuciones de variables de parcela: área, suelo, cultivo, sistema de riego")
bullet(doc, "Cálculo del Índice de Riesgo Hidráulico: IRH = Ky × Kc_ponderado por etapa fenológica")
bullet(doc, "Análisis económico: ahorro proyectado por parcela al alcanzar el KPI de 6,000 m³/ha")
bullet(doc, "Identificación de cultivos con mayor riesgo de estrés hídrico")
bullet(doc, "Correlaciones entre variables agronómicas y consumo de agua histórico")

h2(doc, "1.3 Por qué se usó el Índice de Riesgo Hidráulico (IRH)")
para(doc, (
    "El IRH combina el coeficiente de respuesta al estrés hídrico (Ky, FAO-33) con el coeficiente de cultivo "
    "ponderado (Kc, FAO-56 Tabla 17). La lógica es: cultivos con Ky alto son más sensibles a déficit hídrico, "
    "y el Kc ponderado refleja la demanda evapotranspirativa en la etapa actual del ciclo."
))
para(doc, (
    "Maíz obtiene el IRH más alto (1.125) porque tiene Ky=1.25 (el más alto entre los 5 cultivos) y Kc medio "
    "elevado. Esto justifica priorizar Maíz en las recomendaciones preventivas del sistema."
))

add_table(doc,
    ["Cultivo", "Ky (FAO-33)", "Kc Ponderado", "IRH", "Riesgo"],
    [
        ["Maíz",    "1.25", "0.90", "1.125", "ALTO"],
        ["Algodón",  "0.85", "0.88", "0.748", "MEDIO-ALTO"],
        ["Frijol",  "1.15", "0.65", "0.748", "MEDIO-ALTO"],
        ["Chile",   "1.10", "0.83", "0.913", "ALTO"],
        ["Uva",     "0.85", "0.70", "0.595", "MEDIO"],
    ],
    col_widths_cm=[4, 3.5, 3.5, 3, 4]
)

h2(doc, "1.4 Resultados y Figuras")
para(doc, "Análisis económico clave: ahorro de $3,360 MXN/ha/ciclo representa el diferencial entre la práctica "
         "actual y el objetivo KPI, calculado directamente como (8,000 - 6,000) × $1.68.")

img(doc, IMG_ML, "figura_2___distribuciones_de_variables_de_parcela.png",
    "Figura 1.1 — Distribuciones de variables de parcela (área, tipo de suelo, cultivo, sistema de riego)")
img(doc, IMG_ML, "figura_7___distribuciones_económicas_del_ciclo_agr.png",
    "Figura 1.2 — Distribuciones económicas del ciclo agrícola (costo, ahorro potencial, ROI)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 2 — DATOS SINTÉTICOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "2. Generación y Auditoría de Datos Sintéticos — milpin_datos_sinteticos_auditoria.ipynb")

h2(doc, "2.1 Problema: Sin datos reales disponibles")
para(doc, (
    "El Valle del Yaqui cuenta con registros históricos de riego, pero no están digitalizados en un formato "
    "utilizable, son propiedad de la CONAGUA/módulo, y su acceso requiere convenios institucionales "
    "que exceden el alcance del prototipo. Esta es una restricción real y frecuente en proyectos AgTech en México."
))
para(doc, (
    "La solución adoptada: generar datos sintéticos que sean agronómicamente correctos — calibrados contra "
    "FAO-56, FAO-33, CIMMYT, estadísticas CONAGUA/DR-041 y SIAP — con el objetivo explícito de que el modelo "
    "aprenda relaciones físicas reales, no artefactos estadísticos de la simulación."
))

h2(doc, "2.2 Parámetros del Generador")
info_box(doc, (
    "Parámetros base: seed=42 | 80 parcelas | 20 usuarios | 10 ciclos por parcela | "
    "TARIFA_MXN_M3=1.68 | 5 cultivos (Maíz, Frijol, Algodón, Uva, Chile)"
))

add_table(doc,
    ["Parámetro", "Valor", "Fuente / Justificación"],
    [
        ["Tasa anomalías SOBRE_RIEGO",     "4%",  "Estimación módulo DR-041"],
        ["Tasa AGRICULTOR_INEFICIENTE",    "9%",  "Encuesta piloto Módulo 3"],
        ["Tasa GAP_FALLA_EQUIPO",          "3%",  "Reporte CONAGUA"],
        ["ETo base verano Valle del Yaqui","7.60 mm/día", "NASA POWER + FAO-56 PM"],
        ["Kc Maíz etapa media",            "1.20", "FAO-56 Tabla 17"],
        ["Ky Maíz",                        "1.25", "FAO-33"],
        ["Tarifa agua",                    "$1.68 MXN/m³", "CFE 9-CU, bombeo 80m"],
        ["Salinidad (reducción CE>4 dS/m)","Función sigmoide", "FAO-33 Cap. 3"],
    ],
    col_widths_cm=[5, 3.5, 8]
)

h2(doc, "2.3 Qué valida la Auditoría")
para(doc, "El notebook de auditoría verifica que el generador no inventa — que los datos simulados son "
         "consistentes con las fuentes agronómicas de referencia:")
bullet(doc, "FAO-56: balance hídrico (ETo → ETc → déficit → lamina de riego)")
bullet(doc, "FAO-33: función de respuesta al estrés salino y función Ky para estrés hídrico")
bullet(doc, "CIMMYT: parámetros agronómicos de variedades de Maíz y Trigo en zona semiárida")
bullet(doc, "CONAGUA/DR-041: volúmenes concesionados y distribución histórica de eventos de riego")
bullet(doc, "SIAP: rendimientos esperados por cultivo y municipio en Sonora")

h2(doc, "2.4 Resultados y Figuras")
result_box(doc, "Validación: δ < 5% en todos los parámetros de diseño vs valores generados (concordancia μ)")

img(doc, IMG_ML, "figura_1___parámetros_agronómicos_fao-56_33_por_cu.png",
    "Figura 2.1 — Parámetros agronómicos FAO-56/33 por cultivo (Kc etapas, Ky, profundidad raíz)")
img(doc, IMG_ML, "tabla_2___catálogo_de_cultivos__kc__fao-56_tabla_1.png",
    "Figura 2.2 — Catálogo de cultivos: Kc FAO-56, profundidad de raíz, Ky FAO-33")
img(doc, IMG_ML, "tabla_5___tipos_de_anomalías_inyectadas__ground_tr.png",
    "Figura 2.3 — Tipos de anomalías inyectadas (ground truth para entrenamiento del detector)")
img(doc, IMG_ML, "figura_6___anomalías_inyectadas__impacto_en_distri.png",
    "Figura 2.4 — Impacto de las anomalías inyectadas en la distribución del volumen de riego")
img(doc, IMG_ML, "figura_8___concordancia_μ_diseño_vs_μ_generado__δ.png",
    "Figura 2.5 — Concordancia μ diseño vs μ generado (δ < 5% confirma calibración del generador)")
img(doc, IMG_ML, "tabla_7___resumen_estadístico__parámetros_de_diseñ.png",
    "Figura 2.6 — Resumen estadístico: parámetros de diseño vs valores reales generados")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 3 — FAO-56 VISUALIZACIÓN
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "3. Motor Agronómico FAO-56 — fao56_visualizacion.ipynb")

h2(doc, "3.1 Objetivo")
para(doc, (
    "Este notebook visualiza el motor agronómico central de MILPÍN: la implementación fiel del método FAO-56 "
    "Penman-Monteith (Allen et al. 1998). El propósito es doble: (1) validar que la implementación en "
    "backend/core/balance_hidrico.py produce resultados correctos, y (2) generar figuras de referencia para "
    "el equipo agronómico."
))

h2(doc, "3.2 El Método FAO-56 Penman-Monteith")
para(doc, (
    "Penman-Monteith es el método estándar de la FAO para calcular la evapotranspiración de referencia (ETo). "
    "La ecuación requiere: temperatura máxima y mínima diaria, humedad relativa, velocidad del viento a 2m "
    "y radiación solar. Cuando alguno de estos datos falta, el sistema usa Hargreaves-Samani como fallback, "
    "que solo requiere temperatura y radiación extraterrestre (estimable por latitud y día del año)."
))

add_table(doc,
    ["Método", "Datos requeridos", "RMSE vs PM", "Uso en MILPÍN"],
    [
        ["Penman-Monteith (FAO-56)", "Tmax, Tmin, HR, Viento, Rs", "— (referencia)", "Principal"],
        ["Hargreaves-Samani", "Tmax, Tmin, Ra (latitud/DOY)", "1.25 mm/día", "Fallback"],
    ],
    col_widths_cm=[5, 6, 3.5, 3]
)

h2(doc, "3.3 Resultados Clave")
result_box(doc, (
    "ETo base verano Valle del Yaqui: 7.60 mm/día  |  "
    "ETo anual PM: 1,850 mm/año  |  Hargreaves: 1,657 mm/año  |  "
    "RMSE Hargreaves vs PM: 1.25 mm/día"
))

para(doc, (
    "El ETo de 7.60 mm/día en verano es consistente con la literatura para zonas semiáridas del noroeste "
    "de México. La diferencia anual de ~193 mm entre PM y Hargreaves muestra que el fallback subestima "
    "la demanda evapotranspirativa, lo cual es conservador (recomienda menos riego que el necesario) — "
    "un sesgo aceptable cuando no hay datos meteorológicos completos."
))

h2(doc, "3.4 Figuras")
img(doc, IMG_EXP, "fig1_eto_anual.png",
    "Figura 3.1 — ETo anual calculada con Penman-Monteith (Valle del Yaqui, ciclo tipo)")
img(doc, IMG_EXP, "fig2_pm_vs_hargreaves.png",
    "Figura 3.2 — Comparación PM vs Hargreaves: diferencia diaria y acumulada (RMSE=1.25 mm/día)")
img(doc, IMG_EXP, "fig3_sensibilidad_eto.png",
    "Figura 3.3 — Análisis de sensibilidad: impacto de cada variable en ETo (temperatura > humedad > viento)")
img(doc, IMG_EXP, "fig4_curvas_kc.png",
    "Figura 3.4 — Curvas Kc por cultivo y etapa fenológica (inicial, desarrollo, media, final)")
img(doc, IMG_EXP, "fig8_heatmap_etc.png",
    "Figura 3.5 — Heatmap ETc mensual por cultivo (m³/ha): demanda hídrica total por período")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 4 — BALANCE HÍDRICO
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "4. Balance Hídrico Dinámico — balance_hidrico_visualizacion.ipynb")

h2(doc, "4.1 El Problema que Resuelve")
para(doc, (
    "El sistema original usaba (CC+PMP)/2 como condición inicial de humedad del suelo — básicamente "
    "asumía que toda parcela empieza al 50% de capacidad. Esto es incorrecto: la humedad inicial depende "
    "del último evento de riego, la precipitación, la ETc acumulada y las propiedades del suelo."
))
para(doc, (
    "La función propagar_balance_hidrico() implementada el 2026-05-06 reemplaza ese valor inventado con "
    "un balance acumulado día a día desde el último riego real registrado en la base de datos."
))

h2(doc, "4.2 Lógica de Propagación")
para(doc, "El algoritmo calcula iterativamente para cada día desde el último riego:")
bullet(doc, "Dr(t) = Dr(t-1) + ETc(t) - Pp(t) - Ir(t)   [déficit acumulado, mm]")
bullet(doc, "Si Dr(t) > TAD (Threshold Available Depletion): zona de estrés → recomendación preventiva")
bullet(doc, "Si Dr(t) > RAW (Readily Available Water): riego urgente → lamina de reposición")
bullet(doc, "La propagación es el input principal del endpoint GET /api/parcelas/{id}/balance_hidrico")

h2(doc, "4.3 Impacto en el Sistema")
para(doc, (
    "Con humedad inicial inventada, el sistema generaba recomendaciones que podían diferir hasta 800 m³/ha "
    "del valor correcto en la primera semana tras un riego pesado. La propagación real elimina ese error "
    "sistemático y hace que las recomendaciones sean coherentes con el historial de riego real de cada parcela."
))
result_box(doc, "Solución: propagar_balance_hidrico() — 9 tests unitarios en TestPropagar (51 total en el suite)")

h2(doc, "4.4 Figuras")
img(doc, IMG_EXP, "fig5_balance_hidrico.png",
    "Figura 4.1 — Balance hídrico simulado: déficit (Dr), RAW y TAD por etapa fenológica (Maíz, ciclo PV)")
img(doc, IMG_EXP, "fig6_propagacion.png",
    "Figura 4.2 — Propagación día a día desde último riego: comparación humedad inicial inventada vs real")
img(doc, IMG_EXP, "fig7_costo_riego.png",
    "Figura 4.3 — Costo de riego proyectado: escenario actual vs objetivo KPI por cultivo")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 5 — XGBOOST V3
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "5. XGBoost v3 — milpin_xgboost_prediccion_v3.ipynb")

h2(doc, "5.1 Contexto: La Iteración Anterior")
para(doc, (
    "La versión 3 del modelo XGBoost es la iteración que precedió al pipeline de producción (v4). "
    "Entenderla es importante porque identifica el problema crítico que motivó la versión 4: "
    "la convergencia prematura del modelo por baja diversidad en los datos de entrenamiento."
))

h2(doc, "5.2 Dos Modelos: M1 y M2")
para(doc, "La v3 entrenó dos modelos independientes:")
bullet(doc, "M1: predice volumen_agua_total_m3_ha (regresor directo del KPI principal)")
bullet(doc, "M2: predice rendimiento_real_ton_ha (regresor del rendimiento del cultivo)")
para(doc, (
    "Ambos modelos se entrenan con el mismo pipeline pero features distintas. La separación refleja "
    "la diferencia entre optimizar el uso del agua (M1) y proteger la productividad (M2) — "
    "el sistema debe maximizar ambos simultáneamente."
))

h2(doc, "5.3 El Problema Identificado: Convergencia Prematura")
para(doc, (
    "El generador de datos sintéticos v2 (simple) produce datasets que convergen estadísticamente "
    "a ~20,000-30,000 muestras. Más allá de ese umbral, agregar muestras no mejora el modelo porque "
    "todas las muestras adicionales son estadísticamente indistinguibles de las existentes."
))
para(doc, (
    "Esto es problemático porque el modelo aprende la distribución del generador, no la variabilidad "
    "real del campo. El generador v3 (alta diversidad) resuelve esto introduciendo 5 micro-regiones "
    "geográficas, 4 arquetipos de agricultor, variabilidad ENSO, eventos extremos y ruido IoT — "
    "forzando convergencia real a ~80,000-150,000 muestras."
))

add_table(doc,
    ["Característica", "Generador v2 (simple)", "Generador v3 (alta diversidad)"],
    [
        ["Convergencia estadística",      "~20,000-30,000 muestras", "~80,000-150,000 muestras"],
        ["Micro-regiones geográficas",    "1 (uniforme)",            "5 (suelo, clima diferenciado)"],
        ["Arquetipos de agricultor",      "Ninguno",                 "4 (eficiente, ineficiente, mixto, novato)"],
        ["Variabilidad ENSO",             "No",                      "Sí (año Niño/Niña/neutro)"],
        ["Salinidad FAO-33",              "No",                      "Sí (sigmoide CE vs rendimiento)"],
        ["Ruido IoT en datos climáticos", "No",                      "Sí (sensores defectuosos simulados)"],
        ["Ascenso capilar",               "No",                      "Sí (por tipo de suelo)"],
    ],
    col_widths_cm=[5.5, 4.5, 6.5]
)

h2(doc, "5.4 Por qué v3 → v4")
para(doc, (
    "La v3 validó que el generador de alta diversidad producía modelos más robustos. "
    "La v4 escala esto a 1,000,000 de muestras, agrega optimización bayesiana (Optuna), "
    "pesos de muestra para el rango operacional KPI, y tres submodelos por cultivo en lugar de dos globales."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 6 — XGBOOST V4
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "6. XGBoost v4 — milpin_xgboost_v4.ipynb (Pipeline de Producción)")

h2(doc, "6.1 Objetivo del Experimento")
para(doc, (
    "Este notebook es el pipeline completo de entrenamiento del modelo de producción. El objetivo no es "
    "maximizar R² global — es minimizar el RMSE_op, el error solo dentro del rango operacional "
    "[6,000-10,000 m³/ha] que es donde operan las decisiones de riego reales del sistema. "
    "Un modelo con R²=0.999 pero RMSE_op=200 m³/ha es inútil para el KPI de ahorro del 25%."
))

h2(doc, "6.2 Configuración del Experimento")
add_table(doc,
    ["Parámetro", "Valor", "Justificación"],
    [
        ["N_SAMPLES",       "1,000,000", "Alta diversidad para convergencia real (no estadística)"],
        ["SEED",            "42",        "Reproducibilidad"],
        ["OPTUNA_TRIALS",   "40",        "TPE sampler, equilibrio búsqueda/costo computacional"],
        ["PESO_OP",         "3.0",       "Énfasis 3× en rango operacional KPI"],
        ["RANGO_OP",        "[6,000-10,000] m³/ha", "Rango de decisión real del sistema"],
        ["Split",           "72/8/20%", "Train:720k | Val:80k | Test:200k | HPO:20k/cultivo"],
        ["Modelos/cultivo", "3",         "requiere_riego, lamina_ajustada, riesgo_estres"],
    ],
    col_widths_cm=[4, 4, 8.5]
)

h2(doc, "6.3 Los Tres Submodelos por Cultivo")
para(doc, "Cada uno de los 5 cultivos tiene 3 modelos XGBoost entrenados de forma independiente:")
bullet(doc, "requiere_riego: clasificador binario — ¿regar o esperar? (threshold Optuna-optimizado por cultivo)")
bullet(doc, "lamina_ajustada: regresor de la lámina de riego óptima [mm], con log-transform en target")
bullet(doc, "riesgo_estres: clasificador multiclase — nivel de riesgo de estrés hídrico (bajo/moderado/crítico)")

h2(doc, "6.4 Decisiones de Diseño Clave")

h3(doc, "6.4.1 Log-transform en la lámina de riego")
para(doc, (
    "La variable objetivo (lamina_ajustada) tiene distribución asimétrica con cola derecha — "
    "eventos de riego pesado generan valores extremos. Entrenar sobre log(lamina) reduce la "
    "heteroscedasticidad y evita que el gradiente de los outliers domine el entrenamiento. "
    "Los resultados se desnormalizan con exp() antes de servir la predicción."
))

h3(doc, "6.4.2 Sample weights (PESO_OP=3×)")
para(doc, (
    "Las muestras en el rango [6,000-10,000] m³/ha reciben peso 3.0 vs 1.0 para el resto. "
    "Esto le indica al gradiente de XGBoost que errores en el rango operacional cuestan "
    "3 veces más que errores fuera de él. El resultado es un RMSE_op significativamente mejor "
    "a costa de mayor RMSE global — trade-off deliberado y correcto para el problema de negocio."
))

h3(doc, "6.4.3 Optimización Bayesiana con Optuna (TPE)")
para(doc, (
    "Optuna con el sampler TPE (Tree-structured Parzen Estimator) es más eficiente que grid search "
    "o random search porque modela la distribución de probabilidad de los hiperparámetros buenos "
    "y enfoca las pruebas en regiones prometedoras del espacio de búsqueda. Con 40 trials por cultivo "
    "(5 cultivos × 3 modelos = 15 optimizaciones), el costo computacional es manejable."
))

h3(doc, "6.4.4 58 Features en 4 Generaciones")
para(doc, (
    "Las features se acumularon iterativamente: v2 (9 features base agronómicas) → v3 (+ 7 features "
    "de diversidad) → v4 (+ 10 features de ENSO, salinidad, capilarity, IoT noise). Las 5 features "
    "categóricas (cultivo, región, suelo, arquetipo, año_enso) se encodean con XGBoost nativo "
    "(enable_categorical=True)."
))

h2(doc, "6.5 Advertencia de Leakage de Datos")
para(doc, (
    "Durante el experimento se identificaron dos features con riesgo de leakage de datos. "
    "Este es un hallazgo crítico que afecta la validez del modelo si se ignora:"
))
bullet(doc, "ratio_real_teorico: LEAKAGE SEVERO — el target (vol_real) está en el numerador. "
            "El modelo aprende a dividir el target sobre el teórico, no a predecirlo.")
bullet(doc, "vol_teorico_m3ha: LEAKAGE SUAVE — reconstruye casi perfectamente el target "
            "via balance hídrico FAO-56. En producción esta variable NO está disponible antes del riego.")

info_box(doc,
    "Experimento A/B/C de leakage: Se compararon modelos con ambas features (A), solo vol_teorico (B) "
    "y sin ninguna (C). El modelo sin leakage (C) tiene RMSE_op ~40% mayor pero es el único válido "
    "para producción. Los modelos A y B muestran métricas artificialmente buenas que no se reproducirían "
    "en campo.",
    "FFF3E0"
)

h2(doc, "6.6 Resultados por Cultivo (Optuna HPO)")
add_table(doc,
    ["Cultivo", "RMSE_op (m³/ha)", "Relativo al KPI (2,000 m³/ha)", "Clasificación"],
    [
        ["Chile",   "40 (HPO) / 32.9 (final)", "1.6%", "Mejor"],
        ["Frijol",  "45",                       "2.3%", "Muy bueno"],
        ["Uva",     "53",                       "2.7%", "Bueno"],
        ["Maíz",    "54",                       "2.7%", "Bueno"],
        ["Algodón", "61 (HPO) / 42.0 (final)",  "3.1%", "Aceptable"],
    ],
    col_widths_cm=[4, 5, 5.5, 3]
)

h2(doc, "6.7 Resultados Globales Finales")
result_box(doc, (
    "Global RMSE_op=41.9 m³/ha  |  Ensemble (bagging) RMSE_op=37.8 m³/ha  |  "
    "R²=0.9998  |  Ahorro neto proyectado: 98.1%  |  Chile mejor: RMSE_op=32.9"
))

h2(doc, "6.8 Figuras del Experimento")

img(doc, IMG_EXP, "milpin_dashboard_v4.png",
    "Figura 6.1 — Dashboard principal v4: RMSE_op por cultivo, distribución de errores y métricas globales")
img(doc, IMG_EXP, "rmse_por_rango_comparativa.png",
    "Figura 6.2 — RMSE por rango de volumen: el modelo optimiza el rango operacional [6k-10k] m³/ha")
img(doc, IMG_EXP, "residuales_por_cultivo.png",
    "Figura 6.3 — Residuales por cultivo: distribución de errores (objetivo: centrado en 0, sin sesgo sistemático)")
img(doc, IMG_EXP, "learning_curves_v4.png",
    "Figura 6.4 — Curvas de aprendizaje: train vs validation RMSE_op conforme aumentan las muestras")

h3(doc, "Feature Importance y SHAP")
para(doc, (
    "SHAP (SHapley Additive exPlanations) via el mecanismo nativo de XGBoost (pred_contribs=True) "
    "permite explicar cada predicción individual. Se usa iteration_range para limitar al mejor árbol "
    "según early stopping, evitando sobreajuste en la explicación."
))
img(doc, IMG_EXP, "feature_importance_global.png",
    "Figura 6.5 — Importancia de features global: top 15 por ganancia (gain) en XGBoost")
img(doc, IMG_EXP, "shap_global.png",
    "Figura 6.6 — SHAP global: impacto promedio de cada feature en la magnitud de la predicción")
img(doc, IMG_EXP, "shap_maíz.png",
    "Figura 6.7 — SHAP Maíz: dependencia SHAP vs valor de feature para las 5 más importantes")

h3(doc, "Análisis de Leakage")
img(doc, IMG_EXP, "leakage_metricas_comparativa.png",
    "Figura 6.8 — Comparativa de métricas: modelos A (con leakage), B (suave) y C (sin leakage, producción)")
img(doc, IMG_EXP, "shap_leakage_comparison.png",
    "Figura 6.9 — SHAP leakage: cómo ratio_real_teorico domina la importancia en el modelo contaminado")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 7 — ANOMALY DETECTOR
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "7. Detección de Anomalías — anomaly_detector.ipynb")

h2(doc, "7.1 Objetivo")
para(doc, (
    "El detector de anomalías identifica parcelas con patrones de riego anómalos: sobre-riego, "
    "agricultor ineficiente sistemático, o fallas de equipo (gaps prolongados sin riego). "
    "Es un modelo no supervisado porque en campo real no existe un label 'anómalo' — el ground truth "
    "son las etiquetas inyectadas por el generador sintético, usadas solo para evaluación."
))
para(doc, (
    "El nivel de análisis es (parcela × ciclo), no eventos individuales. Esto evita que un solo "
    "evento de riego pesado dispare una alerta falsa — lo que importa es el patrón agregado del ciclo."
))

h2(doc, "7.2 Pipeline de Datos")
bullet(doc, "Datos crudos: 13,427 eventos de riego individuales")
bullet(doc, "Agregación: → 960 pares (parcela × ciclo) con 7 features estadísticas")
bullet(doc, "Features: vol_total_m3_ha, vol_media_evento, vol_cv (coeficiente de variación), "
            "n_eventos, max_gap_dias, costo_total_mxn, sistema_riego_num")
bullet(doc, "Modelo: Isolation Forest (contamination=0.12, n_estimators=160, random_state=42)")

h2(doc, "7.3 Por qué Isolation Forest")
para(doc, (
    "Isolation Forest es adecuado para este problema porque: (1) no requiere labels de entrenamiento, "
    "(2) escala bien a miles de parcelas, (3) el parámetro contamination permite calibrar la tasa "
    "esperada de anomalías (12% = suma de las tres tasas del generador: 4+9+3%), y (4) produce "
    "una puntuación continua (anomaly score) que permite rankear parcelas por nivel de riesgo."
))

h2(doc, "7.4 Limitación: Sin Feature Importance Nativa")
para(doc, (
    "A diferencia de XGBoost, Isolation Forest no produce feature importances directamente. "
    "La solución adoptada es Cohen's d: se calcula el tamaño del efecto estandarizado entre la "
    "distribución de cada feature en parcelas anómalas vs normales. Un Cohen's d alto indica que "
    "esa feature discrimina bien entre los dos grupos."
))
para(doc, (
    "La interpretación de Cohen's d: trivial (<0.2), pequeño (0.2-0.5), mediano (0.5-0.8), grande (>0.8). "
    "Para acciones operacionales, solo features con d>0.5 (mediano o grande) justifican intervención."
))

h2(doc, "7.5 Resultados")
add_table(doc,
    ["Feature", "Cohen's d", "Magnitud", "Interpretación"],
    [
        ["sistema_riego_num", "0.829", "Grande",  "Tipo de riego es el predictor más fuerte de anomalía"],
        ["vol_cv",            "0.634", "Mediano", "Alta variabilidad entre eventos → patrón irregular"],
        ["costo_total_mxn",   "0.592", "Mediano", "Costo total refleja sobre-riego monetario"],
        ["n_eventos",         "0.518", "Mediano", "Frecuencia anormal de eventos de riego"],
        ["vol_total_m3_ha",   "0.472", "Pequeño", "Volumen total: discrimina pero con superposición"],
        ["vol_media_evento",  "0.333", "Pequeño", "Lámina por evento: menos diagnóstica"],
        ["max_gap_dias",      "0.180", "Trivial", "Gap máximo: poco discriminativo globalmente"],
    ],
    col_widths_cm=[4.5, 2.5, 3, 6.5]
)

result_box(doc, (
    "Isolation Forest — Resultados: Precision=0.586  |  Recall=0.466  |  F1=0.519  |  "
    "Contamination=0.12  |  N=960 pares parcela×ciclo"
))

h2(doc, "7.6 Recall por Tipo de Anomalía")
add_table(doc,
    ["Tipo de Anomalía", "Recall", "Interpretación"],
    [
        ["GAP_FALLA_EQUIPO",      "57.1%", "El mejor — gaps largos generan scores de aislamiento altos"],
        ["AGRICULTOR_INEFICIENTE","45.0%", "Aceptable — patrón sistemático detectable en el ciclo"],
        ["SOBRE_RIEGO",           "42.1%", "Más difícil — eventos individuales se diluyen en el promedio"],
    ],
    col_widths_cm=[5, 3, 8.5]
)

para(doc, (
    "Un Recall de 0.466 en un problema no supervisado es un resultado razonable, no deficiente. "
    "El baseline teórico de un detector aleatorio con contamination=0.12 tiene Recall=0.12 "
    "(solo detecta lo que etiqueta al azar). El modelo mejora este baseline en 3.9×."
))

h2(doc, "7.7 Caso Extremo Identificado")
info_box(doc,
    "Top anomalía: Parcela OI-2024  |  Anomaly Score=-0.705  |  Vol=16,829 m³/ha  |  "
    "Label: AGRICULTOR_INEFICIENTE  |  Excede 2.1× el KPI de 8,000 m³/ha baseline",
    "FFE8E8"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 8 — ARQUITECTURA ML EN PRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "8. Arquitectura de Inferencia y MLOps")

h2(doc, "8.1 Separación Training / Inference")
para(doc, (
    "La regla central de la arquitectura: ml/training/ nunca importa backend/. "
    "El código de entrenamiento es independiente del sistema de producción. "
    "Los modelos se exportan como artefactos (archivos .pkl/.json) que ml/inference/ carga en runtime."
))
para(doc, (
    "Esta separación evita el problema clásico de 'train/serve skew': que el código de preprocessing "
    "en entrenamiento difiera del código en inferencia. El módulo ml/inference/feature_preprocessor.py "
    "es la única fuente de verdad para la transformación de features."
))

h2(doc, "8.2 Singleton Pattern en Inferencia")
para(doc, (
    "El wrapper ml/inference/xgboost_riego.py implementa el patrón Singleton via obtener_predictor(). "
    "Los 15 modelos (5 cultivos × 3 submodelos) se cargan una sola vez al primer request y permanecen "
    "en memoria. Las cargas subsiguientes retornan la instancia existente."
))
bullet(doc, "11 FEATURE_NAMES en orden fijo — el orden importa para XGBoost en producción")
bullet(doc, "PrediccionRiego dataclass: requiere_riego, lamina_ajustada, riesgo_estres, urgencia, score_confianza")
bullet(doc, "Fallback FAO-56: si los modelos no están disponibles, el sistema usa balance hídrico agronómico")
bullet(doc, "_clasificar_urgencia(): 'crítico' | 'moderado' | 'preventivo' según déficit y Kc")

h2(doc, "8.3 Forecast de ETo a 7 Días")
para(doc, (
    "El módulo backend/core/eto_forecast.py usa Ridge Regression sobre clima_diario con features: "
    "sin/cos del día del año (captura estacionalidad), lags de ETo (t-1, t-3, t-7), temperatura máxima. "
    "Si hay menos de 60 registros históricos, usa la media de los últimos 14 días como fallback."
))
para(doc, (
    "El endpoint GET /api/parcelas/{id}/forecast?dias_siembra=N&horizon=7 proyecta ETo 7 días "
    "y corre FAO-56 forward para estimar la fecha del próximo riego con intervalo de incertidumbre (±días)."
))

h2(doc, "8.4 Feature Store")
para(doc, "El Feature Store define las features en YAML para separar la lógica de materialización de la lógica de negocio:")
bullet(doc, "parcela_static.yaml: área, tipo_suelo, sistema_riego, coordenadas — estático por ciclo")
bullet(doc, "parcela_daily.yaml: ETo, Kc, Dr, precipitación — actualización diaria desde clima_diario")
bullet(doc, "parcela_ciclo.yaml: features agregadas del ciclo actual — rol_cumplimiento_kpi, eficiencia")

h2(doc, "8.5 Monitoreo de Drift")

add_table(doc,
    ["Métrica", "Umbral", "Acción"],
    [
        ["PSI (Population Stability Index)", "< 0.1 → estable", "Sin acción"],
        ["PSI", "0.1 – 0.2 → monitorear",  "Alerta amarilla, revisar próxima semana"],
        ["PSI", "> 0.2 → retrain",          "Alerta roja, disparar pipeline de reentrenamiento"],
        ["KS test p-value", "< 0.05 → drift detectado", "Investigar feature, posible retrain"],
    ],
    col_widths_cm=[5.5, 4, 7]
)

para(doc, (
    "drift.py implementa calcular_psi() y calcular_ks() sobre las distribuciones de features "
    "de producción vs entrenamiento. Los runbooks en docs/runbooks/ describen el procedimiento "
    "de respuesta a cada tipo de alerta."
))

h2(doc, "8.6 Promote Gate")
para(doc, (
    "ml/training/xgboost_riego/promote.py implementa una compuerta de promoción automática: "
    "el modelo candidato solo reemplaza al de producción si supera umbrales definidos en "
    "ml/configs/xgboost_riego.yaml. Los umbrales actuales: RMSE_op < 50 m³/ha por cultivo, "
    "F1 clasificador riesgo > 0.70. Esto previene que un reentrenamiento con datos drift "
    "degrade el modelo de producción."
))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 9 — RESULTADOS CONSOLIDADOS
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "9. Resultados Consolidados y Validación")

h2(doc, "9.1 Resumen de Modelos en Producción")
add_table(doc,
    ["Modelo", "Tipo", "N Entrenamiento", "Métrica Principal", "Resultado"],
    [
        ["XGBoost v4 (5 cultivos × 3)",    "Regresión / Clasif.", "720,000", "RMSE_op [6k-10k]", "37.8 m³/ha (ensemble)"],
        ["Isolation Forest",                "Anomalía (no superv.)","960 pares ciclo","F1",          "0.519"],
        ["Ridge Regression ETo",            "Serie temporal",      "≥60 días",      "RMSE ETo 7d",  "Fallback: media 14d"],
        ["FAO-56 Penman-Monteith",          "Motor físico",        "N/A",           "RMSE vs ref.", "0 mm/día (referencia)"],
        ["Hargreaves-Samani (fallback)",    "Motor físico",        "N/A",           "RMSE vs PM",   "1.25 mm/día"],
    ],
    col_widths_cm=[5, 3.5, 3.5, 3.5, 4]
)

h2(doc, "9.2 Impacto en el KPI")
para(doc, (
    "El RMSE_op global de 37.8 m³/ha sobre un KPI de ahorro de 2,000 m³/ha representa un error "
    "relativo del 1.9%. En términos prácticos: si el sistema recomienda regar con 6,200 m³/ha, "
    "la incertidumbre de la predicción es ±38 m³/ha — precisión más que suficiente para tomar "
    "decisiones de riego en campo."
))

result_box(doc, (
    "Error relativo sobre KPI: 1.9%  |  RMSE_op=37.8 m³/ha  vs  Δ_KPI=2,000 m³/ha  |  "
    "Ahorro proyectado: $3,360 MXN/ha/ciclo  |  Precisión operacional: suficiente para decisiones de campo"
))

h2(doc, "9.3 Deuda Técnica ML Conocida")
bullet(doc, "Re-exports temporales en backend/core/: eliminar en Fase B y usar ml/inference/ directamente")
bullet(doc, "Pipelines Prefect: actualmente stubs — automatización real es Fase C/D")
bullet(doc, "MLflow registry: docs/MLOPS.md documenta la integración planeada pero no implementada")
bullet(doc, "Tests ML: test_preprocessor.py, test_drift.py, test_promote_gate.py son stubs en ml/tests/")
bullet(doc, "Isolation Forest: Recall=0.466 es mejorable con features de secuencia temporal (LSTM, señales GIS)")

h2(doc, "9.4 Próximos Pasos Priorizados")
add_table(doc,
    ["Fase", "Tarea", "Impacto", "Esfuerzo"],
    [
        ["A (Higiene)", "Rotar credenciales BD y sanitizar path traversal en voice_endpoint.py", "Alto (seguridad)", "Bajo"],
        ["A", "Mover catálogo cultivos a tabla BD (eliminar constantes duplicadas en 6 archivos)", "Medio", "Medio"],
        ["B (Features)", "Completar ml/tests/ con tests reales de preprocessor y drift", "Alto (confiabilidad)", "Medio"],
        ["B", "Eliminar re-exports temporales backend/core/xgboost_riego.py y anomaly_detector.py", "Bajo", "Bajo"],
        ["C (MLOps)", "Activar pipelines Prefect: nasa_power_daily + train_eval_promote", "Alto (autonomía)", "Alto"],
        ["C", "Conectar MLflow registry con promote gate", "Medio", "Medio"],
        ["D (Datos reales)", "Integrar datos históricos reales DR-041 cuando estén disponibles", "Crítico", "Alto"],
    ],
    col_widths_cm=[2.5, 7, 3.5, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# CAPÍTULO 10 — CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════
h1(doc, "10. Conclusiones")

para(doc, (
    "El módulo ML de MILPÍN demuestra que es técnicamente viable construir un sistema de recomendación "
    "de riego preciso usando datos sintéticos calibrados contra estándares agronómicos internacionales "
    "(FAO-56, FAO-33, CIMMYT). La clave no fue el algoritmo — XGBoost es una elección pragmática — "
    "sino el rigor en la generación de datos y la definición de la métrica correcta (RMSE_op en el "
    "rango operacional, no R² global)."
))

para(doc, (
    "Los hallazgos más importantes del experimento son: (1) el leakage de datos es la mayor amenaza "
    "para modelos agrícolas donde las variables físicas se interrelacionan; (2) la diversidad del "
    "generador sintético importa más que el volumen de datos; y (3) el motor físico FAO-56 es más "
    "valioso como fallback y como fuente de features que como competidor del modelo ML."
))

para(doc, (
    "El siguiente paso crítico no es mejorar el modelo — es validarlo con datos reales del DR-041 "
    "y exponer el sistema a agricultores reales en el Módulo 3. Un RMSE_op de 37.8 m³/ha con datos "
    "sintéticos podría degradarse o mejorar con datos reales; la única forma de saberlo es medir."
))

result_box(doc, (
    "Estado final: Pre-MVP con core ML funcional  |  "
    "37.8 m³/ha RMSE_op (ensemble)  |  F1=0.519 anomalías  |  "
    "Próximo hito: validación con datos reales DR-041"
))

# ══════════════════════════════════════════════════════════════════════════
# GUARDAR
# ══════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"✅ Documento generado: {OUT}")
print(f"   Tamaño: {os.path.getsize(OUT):,} bytes")
